from imports import (
    asyncio,
    json,
    os,
    fitz,
    requests,
    types
)
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
from config import bot, api_key, log
from bots_dicts import *

# Словарь для хранения текущего режима пользователя
current_mode = {}

# Словари для хранения данных пользователя в режиме сопроводительного письма
resume = {}
profession = {}
company = {}
description = {}

def return_to_main_menu(message):
    """Возвращает пользователя в главное меню.
    
    Args:
        message (types.Message): Объект сообщения от пользователя
    """
    log.info(f"User {message.from_user.id} returned to main menu")
    
    # Очищаем текущий режим
    if message.from_user.id in current_mode:
        del current_mode[message.from_user.id]
        
    welcome_text = (
        "👋 Привет! Я главный бот YourOffer.\n\n"
        "Я помогу тебе с:\n"
        "📝 Написанием сопроводительного письма\n"
        "📄 Созданием резюме\n"
        "🤖 Подготовкой к собеседованию\n"
        "🔍 Поиском вакансий\n\n"
        "Выбери нужный режим работы:"
    )
    bot.send_message(
        message.chat.id,
        welcome_text,
        reply_markup=create_main_menu()
    )


def create_main_menu():
    """Создает главное меню бота с основными кнопками.

    Returns:
        types.ReplyKeyboardMarkup: Объект клавиатуры с кнопками меню
    """
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    buttons = [
        types.KeyboardButton("📝 Сопроводительное письмо"),
        types.KeyboardButton("📄 Резюме"),
        types.KeyboardButton("🤖 AI Интервьюер"),
        types.KeyboardButton("🔍 Парсер вакансий")
    ]
    markup.add(*buttons)
    return markup


def create_main_menu_button():
    """Создает меню только с кнопкой главного меню.

    Returns:
        types.ReplyKeyboardMarkup: Объект клавиатуры с кнопкой главного меню
    """
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    markup.add(types.KeyboardButton("🏠 Главное меню"))
    return markup


def add_main_menu_button(message):
    """Добавляет кнопку возврата в главное меню.
    
    Args:
        message (types.Message): Объект сообщения от пользователя
    """
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    markup.add(types.KeyboardButton("🏠 Главное меню"))
    bot.send_message(message.chat.id, "Нажмите кнопку ниже, чтобы вернуться в главное меню:", reply_markup=markup)


# Функции для cover_letter_bot
def async_handler(f):
    """Декоратор для асинхронных обработчиков сообщений.

    Args:
        f (function): Асинхронная функция-обработчик

    Returns:
        function: Обертка для асинхронной функции
    """

    def wrapper(*args):
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        loop.run_until_complete(f(*args))

    return wrapper


def process_pdf(file_path: str) -> str:
    """Извлекает текст из PDF файла.

    Args:
        file_path (str): Путь к PDF файлу

    Returns:
        str: Извлеченный текст из PDF
    """
    doc = fitz.open(file_path)
    text = ""
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text += page.get_text()
    return text


async def send_prompt_to_gpt(prompt):
    """Отправляет запрос к GPT API и получает ответ.

    Args:
        prompt (str): Текст запроса к GPT

    Returns:
        str: Ответ от GPT или None в случае ошибки
    """
    endpoint = 'https://api.openai.com/v1/chat/completions'

    headers = {
        'Content-Type': 'application/json',
        'Authorization': f'Bearer {api_key}'
    }

    data = {
        'model': 'gpt-4o-mini',
        'messages': [{'role': 'user', 'content': prompt}],
        'max_tokens': 3000,
        'top_p': 1.0,
        'temperature': 0.6
    }

    response = requests.post(endpoint, headers=headers, json=data)

    if response.status_code == 200:
        response_data = json.loads(response.text)
        return response_data['choices'][0]['message']['content']
    else:
        print("Error:", response.text)
        return None


def cover_letter_start(message):
    """Начинает процесс создания сопроводительного письма.

    Args:
        message (types.Message): Объект сообщения от пользователя
    """
    log.info(f"Starting cover letter bot for user {message.from_user.id}")
    
    # Сохраняем текущий режим
    current_mode[message.from_user.id] = "cover_letter"
    
    # Добавляем только кнопку главного меню при старте
    markup = create_main_menu_button()
    bot.send_message(
        message.chat.id,
        "Привет! Я помогу тебе написать сопроводительное письмо.\n"
        "Отправь, пожалуйста, свое резюме в текстовом формате, или в форматах pdf или doc",
        reply_markup=markup
    )
    bot.register_next_step_handler(message, ask_resume_async)


async def ask_resume(message):
    """Запрашивает резюме у пользователя и обрабатывает его.

    Args:
        message (types.Message): Объект сообщения от пользователя
    """
    if message.text == "🏠 Главное меню":
        return_to_main_menu(message)
        return

    if message.content_type == 'text':
        resume[message.chat.id] = message.text
    elif message.content_type == 'document' and (
            message.document.file_name.lower().endswith('.pdf') or
            message.document.file_name.lower().endswith('.doc')
    ):
        try:
            file_info = bot.get_file(message.document.file_id)
            downloaded_file = bot.download_file(file_info.file_path)

            local_file_path = os.path.join(
                "Documents",
                file_info.file_path.split('/')[-1]
            )
            os.makedirs(os.path.dirname(local_file_path), exist_ok=True)

            with open(local_file_path, 'wb') as new_file:
                new_file.write(downloaded_file)

            resume[message.chat.id] = process_pdf(local_file_path)
            os.remove(local_file_path)
        except Exception as e:
            log.error(f"Error processing document: {e}")
            bot.send_message(
                message.chat.id,
                "Произошла ошибка :( Пожалуйста, вернитесь в главное меню"
            )
            return_to_main_menu(message)
            return
    else:
        bot.send_message(
            message.chat.id,
            "Пожалуйста, отправьте резюме в текстовом формате или в форматах PDF/DOC"
        )
        bot.register_next_step_handler(message, ask_resume_async)
        return

    bot.send_message(
        message.chat.id,
        "Введите название искомой профессии:"
    )
    bot.register_next_step_handler(message, ask_profession)


ask_resume_async = async_handler(ask_resume)


def ask_profession(message):
    """Запрашивает название профессии у пользователя.

    Args:
        message (types.Message): Объект сообщения от пользователя
    """
    if message.text == "🏠 Главное меню":
        return_to_main_menu(message)
        return

    if message.content_type != 'text':
        bot.send_message(
            message.chat.id,
            "Пожалуйста, введите название профессии текстом."
        )
        bot.register_next_step_handler(message, ask_profession)
        return

    profession[message.chat.id] = message.text
    bot.send_message(
        message.chat.id,
        "Введите название компании, в которую хотите устроиться:"
    )
    bot.register_next_step_handler(message, ask_company)


def ask_company(message):
    """Запрашивает название компании у пользователя.

    Args:
        message (types.Message): Объект сообщения от пользователя
    """
    if message.text == "🏠 Главное меню":
        return_to_main_menu(message)
        return

    if message.content_type != 'text':
        bot.send_message(
            message.chat.id,
            "Пожалуйста, введите название компании текстом."
        )
        bot.register_next_step_handler(message, ask_company)
        return

    company[message.chat.id] = message.text
    bot.send_message(
        message.chat.id,
        "Расскажите о себе в 2-3 предложениях:"
    )
    bot.register_next_step_handler(message, ask_description)


def ask_description(message):
    """Запрашивает описание пользователя.

    Args:
        message (types.Message): Объект сообщения от пользователя
    """
    if message.text == "🏠 Главное меню":
        return_to_main_menu(message)
        return

    if message.content_type != 'text':
        bot.send_message(
            message.chat.id,
            "Пожалуйста, введите описание текстом."
        )
        bot.register_next_step_handler(message, ask_description)
        return

    description[message.chat.id] = message.text
    
    # Генерируем сопроводительное письмо
    prompt = f"""Напиши сопроводительное письмо для соискателя на должность {profession[message.chat.id]} в компанию {company[message.chat.id]}.
    
    Резюме соискателя:
    {resume[message.chat.id]}
    
    Описание соискателя:
    {description[message.chat.id]}
    
    Письмо должно быть профессиональным, но не слишком формальным. Включи описание пользователя из последнего запроса.
    """
    
    cover_letter = send_prompt_to_gpt_sync(prompt)
    
    # Отправляем сопроводительное письмо
    bot.send_message(
        message.chat.id,
        f"Вот ваше сопроводительное письмо:\n\n{cover_letter}"
    )
    
    # Добавляем кнопки рестарта и главного меню
    markup = create_restart_menu()
    bot.send_message(
        message.chat.id,
        "Выберите действие:",
        reply_markup=markup
    )


def restart_cover_letter(message):
    """Перезапускает режим создания сопроводительного письма."""
    try:
        log.info(f"Пользователь {message.from_user.id} начал перезапуск режима сопроводительного письма")
        
        # Очищаем все данные пользователя
        log.info(f"Очистка данных пользователя {message.from_user.id} в режиме сопроводительного письма")
        if message.chat.id in resume:
            del resume[message.chat.id]
        if message.chat.id in profession:
            del profession[message.chat.id]
        if message.chat.id in company:
            del company[message.chat.id]
        if message.chat.id in description:
            del description[message.chat.id]
        
        # Перезапускаем режим с сохранением текущего режима
        current_mode[message.from_user.id] = "cover_letter"
        log.info(f"Установлен режим 'cover_letter' для пользователя {message.from_user.id}")
        
        # Отправляем приветственное сообщение с кнопкой главного меню
        markup = create_main_menu_button()
        log.info(f"Отправка приветственного сообщения пользователю {message.from_user.id}")
        bot.send_message(
            message.chat.id,
            "Давайте начнем заново!\n\n"
            "Отправь, пожалуйста, свое резюме в текстовом формате, или в форматах pdf или doc",
            reply_markup=markup
        )
        bot.register_next_step_handler(message, ask_resume_async)
        log.info(f"Установлен обработчик следующего шага для пользователя {message.from_user.id}")
        
    except Exception as e:
        log.error(f"Ошибка при перезапуске режима сопроводительного письма для пользователя {message.from_user.id}: {str(e)}")
        # В случае ошибки возвращаем в главное меню
        return_to_main_menu(message)


# Функции для resume_bot
def resume_bot_start(message):
    """Начинает процесс создания резюме.

    Args:
        message (types.Message): Объект сообщения от пользователя
    """
    log.info(f"Starting resume bot for user {message.from_user.id}")

    # Сохраняем текущий режим
    current_mode[message.from_user.id] = "resume"
    
    # Добавляем только кнопку главного меню при старте
    markup = create_main_menu_button()
    bot.send_message(
        message.chat.id,
        "Привет! Я помогу тебе составить резюме. Давай для начала познакомимся. "
        "Напиши, пожалуйста, свое ФИО.",
        reply_markup=markup
    )
    
    answers_X[message.chat.id] = ''
    answers_Y[message.chat.id] = ''
    answers_Z[message.chat.id] = ''
    dialogue[message.chat.id] = ''
    projects[message.chat.id] = []
    context[message.chat.id] = []
    question_counter[message.chat.id] = 1
    bot.register_next_step_handler(message, user_name)


def user_name(message):
    """Обрабатывает ввод имени пользователя для создания резюме.

    Args:
        message (types.Message): Объект сообщения от пользователя
    """
    if message.text == "🏠 Главное меню":
        return_to_main_menu(message)
        return

    if message.content_type != 'text':
        bot.send_message(
            message.chat.id,
            "Пожалуйста, отправь свое ФИО текстом."
        )
        bot.register_next_step_handler(message, user_name)
        return
    else:
        name[message.chat.id] = message.text
        bot.send_message(
            message.chat.id,
            "Расскажи о себе в двух-трех предложениях."
        )
        bot.register_next_step_handler(message, user_summary_async)
        return


async def user_summary(message):
    """Обрабатывает ввод краткого описания пользователя.

    Args:
        message (types.Message): Объект сообщения от пользователя
    """
    if message.text == "🏠 Главное меню":
        return_to_main_menu(message)
        return

    if message.content_type not in ('text', 'voice'):
        bot.send_message(
            message.chat.id,
            "Пожалуйста, отправь либо текст, либо голосовое сообщение."
        )
        bot.register_next_step_handler(message, user_summary_async)
        return
    else:
        summary[message.chat.id] = message.text
        bot.send_message(
            message.chat.id,
            "Расскажи о каком-нибудь своем проекте. Опиши его и расскажи, "
            "чем ты в нем занимался."
        )
        dialogue[message.chat.id] = (
            "Вопрос №1: Расскажи о каком-нибудь своем проекте. Опиши его и "
            "расскажи, чем ты в нем занимался."
        )
        answers_X[message.chat.id] = (
            "Вопрос №1: Расскажи о каком-нибудь своем проекте. Опиши его и "
            "расскажи, чем ты в нем занимался."
        )
        bot.register_next_step_handler(message, ask_questions_X_async)
        return


user_summary_async = async_handler(user_summary)


async def ask_questions_X(message):
    """Задает вопросы о проектах пользователя и обрабатывает ответы.

    Args:
        message (types.Message): Объект сообщения от пользователя
    """
    if message.text == "🏠 Главное меню":
        return_to_main_menu(message)
        return

    if message.content_type not in ('text', 'voice'):
        bot.send_message(
            message.chat.id,
            "Пожалуйста, отправь либо текст, либо голосовое сообщение."
        )
        bot.register_next_step_handler(message, ask_questions_X_async)
        return
    else:
        text = message.text
        dialogue[message.chat.id] += f"\nОтвет: {text}\n\n"
        answers_X[message.chat.id] += f"\nОтвет: {text}\n\n"

        grade = await completeness(
            "Оцени числом от 1 до 10 насколько полно я ответил на первоначальный вопрос.",
            answers_X[message.chat.id],
            message.chat.id
        )
        try:
            # Извлекаем число из ответа GPT
            import re
            numbers = re.findall(r'\d+', grade)
            if numbers:
                grade = int(numbers[0])
            else:
                grade = 5  # Значение по умолчанию, если число не найдено
        except (ValueError, IndexError):
            grade = 5  # Значение по умолчанию при любой ошибке

        if grade > 5:
            follow_up[message.chat.id] = await ask_follow_up(
                "Придумай дополнительный вопрос, который бы лучше раскрывал "
                "мой ответ на первоначальный вопрос.",
                answers_X[message.chat.id],
                '',
                message.chat.id
            )
            question_counter[message.chat.id] += 1

            sent_message = bot.send_message(
                message.chat.id,
                follow_up[message.chat.id]
            )
            dialogue[message.chat.id] += (
                f"Вопрос №{question_counter[message.chat.id]}: "
                f"{sent_message.text}"
            )
            answers_X[message.chat.id] += (
                f"Вопрос №{question_counter[message.chat.id]}: "
                f"{sent_message.text}"
            )

            bot.register_next_step_handler(message, ask_questions_X_async)
        else:
            question_counter[message.chat.id] += 1
            sent_message = bot.send_message(
                message.chat.id,
                "Какие инструменты ты использовал при реализации этого проекта?"
            )

            dialogue[message.chat.id] += (
                f"Вопрос №{question_counter[message.chat.id]}: "
                f"{sent_message.text}"
            )
            answers_Y[message.chat.id] += (
                f"Вопрос №{question_counter[message.chat.id]}: "
                f"{sent_message.text}"
            )

            bot.register_next_step_handler(message, ask_questions_Y_async)


ask_questions_X_async = async_handler(ask_questions_X)


async def ask_questions_Y(message):
    """Задает вопросы об инструментах, использованных в проекте.

    Args:
        message (types.Message): Объект сообщения от пользователя
    """
    if message.text == "🏠 Главное меню":
        return_to_main_menu(message)
        return

    if message.content_type not in ('text', 'voice'):
        bot.send_message(
            message.chat.id,
            "Пожалуйста, отправь либо текст, либо голосовое сообщение"
        )
        bot.register_next_step_handler(message, ask_questions_Y_async)
    else:
        text = message.text
        dialogue[message.chat.id] += f"\nОтвет: {text}\n\n"
        answers_Y[message.chat.id] += f"\nОтвет: {text}\n\n"

        grade = await completeness(
            "Оцени числом от 1 до 10 насколько полно я ответил на первоначальный вопрос.",
            answers_Y[message.chat.id],
            message.chat.id
        )
        try:
            # Извлекаем число из ответа GPT
            import re
            numbers = re.findall(r'\d+', grade)
            if numbers:
                grade = int(numbers[0])
            else:
                grade = 5  # Значение по умолчанию, если число не найдено
        except (ValueError, IndexError):
            grade = 5  # Значение по умолчанию при любой ошибке

        if grade > 5:
            follow_up[message.chat.id] = await ask_follow_up(
                "Придумай дополнительный вопрос, который бы лучше раскрывал "
                "мой ответ на первоначальный вопрос.",
                answers_Y[message.chat.id],
                '',
                message.chat.id
            )
            question_counter[message.chat.id] += 1

            sent_message = bot.send_message(
                message.chat.id,
                follow_up[message.chat.id]
            )
            dialogue[message.chat.id] += (
                f"Вопрос №{question_counter[message.chat.id]}: "
                f"{sent_message.text}"
            )
            answers_Y[message.chat.id] += (
                f"Вопрос №{question_counter[message.chat.id]}: "
                f"{sent_message.text}"
            )

            bot.register_next_step_handler(message, ask_questions_Y_async)
        else:
            question_counter[message.chat.id] += 1
            sent_message = bot.send_message(
                message.chat.id,
                "К чему привел этот проект? Можно ли как-то измерить степень его успешности?"
            )

            dialogue[message.chat.id] += (
                f"Вопрос №{question_counter[message.chat.id]}: "
                f"{sent_message.text}"
            )
            answers_Y[message.chat.id] += (
                f"Вопрос №{question_counter[message.chat.id]}: "
                f"{sent_message.text}"
            )

            bot.register_next_step_handler(message, ask_questions_Z_async)


ask_questions_Y_async = async_handler(ask_questions_Y)


async def ask_questions_Z(message):
    """Задает вопросы о результатах проекта.

    Args:
        message (types.Message): Объект сообщения от пользователя
    """
    if message.text == "🏠 Главное меню":
        return_to_main_menu(message)
        return

    if message.content_type not in ('text', 'voice'):
        bot.send_message(
            message.chat.id,
            "Пожалуйста, отправь либо текст, либо голосовое сообщение"
        )
        bot.register_next_step_handler(message, ask_questions_Z_async)
    else:
        text = message.text
        dialogue[message.chat.id] += f"\nОтвет: {text}\n\n"
        answers_Z[message.chat.id] += f"\nОтвет: {text}\n\n"

        grade = await completeness(
            "Оцени числом от 1 до 10 насколько полно я ответил на первоначальный вопрос.",
            answers_Y[message.chat.id],
            message.chat.id
        )
        try:
            # Извлекаем число из ответа GPT
            import re
            numbers = re.findall(r'\d+', grade)
            if numbers:
                grade = int(numbers[0])
            else:
                grade = 5  # Значение по умолчанию, если число не найдено
        except (ValueError, IndexError):
            grade = 5  # Значение по умолчанию при любой ошибке

        if grade > 5:
            follow_up[message.chat.id] = await ask_follow_up(
                "Придумай дополнительный вопрос, который бы лучше раскрывал "
                "мой ответ на первоначальный вопрос.",
                answers_Z[message.chat.id],
                '',
                message.chat.id
            )
            question_counter[message.chat.id] += 1

            sent_message = bot.send_message(
                message.chat.id,
                follow_up[message.chat.id]
            )
            dialogue[message.chat.id] += (
                f"Вопрос №{question_counter[message.chat.id]}: "
                f"{sent_message.text}"
            )
            answers_Z[message.chat.id] += (
                f"Вопрос №{question_counter[message.chat.id]}: "
                f"{sent_message.text}"
            )

            bot.register_next_step_handler(message, ask_questions_Z_async)
        else:
            markup = types.InlineKeyboardMarkup()
            markup.add(
                types.InlineKeyboardButton(
                    'Да',
                    callback_data=f'да\n{message.chat.id}'
                )
            )
            markup.add(
                types.InlineKeyboardButton(
                    'Нет',
                    callback_data=f'нет\n{message.chat.id}'
                )
            )

            projects[message.chat.id].append(dialogue[message.chat.id])
            sent_message = bot.send_message(
                message.chat.id,
                "Отлично! Спасибо за твои ответы. Хочешь рассказать о каком-нибудь "
                "еще из своих проектов?",
                reply_markup=markup
            )
            previous_message_id[message.chat.id] = sent_message.message_id


ask_questions_Z_async = async_handler(ask_questions_Z)


async def user_achievements(message):
    """Обрабатывает ввод достижений пользователя.

    Args:
        message (types.Message): Объект сообщения от пользователя
    """
    if message.text == "🏠 Главное меню":
        return_to_main_menu(message)
        return

    if message.content_type not in ('text', 'voice'):
        bot.send_message(
            message.chat.id,
            "Пожалуйста, отправь либо текст, либо голосовое сообщение"
        )
        bot.register_next_step_handler(message, user_achievements_async)
    else:
        achievements[message.chat.id] = message.text
        bot.send_message(
            message.chat.id,
            "Какими навыками ты обладаешь?"
        )
        bot.register_next_step_handler(message, user_skills_async)
    return


user_achievements_async = async_handler(user_achievements)


async def user_skills(message):
    """Обрабатывает ввод навыков пользователя.

    Args:
        message (types.Message): Объект сообщения от пользователя
    """
    if message.text == "🏠 Главное меню":
        return_to_main_menu(message)
        return

    if message.content_type not in ('text', 'voice'):
        bot.send_message(
            message.chat.id,
            "Пожалуйста, отправь либо текст, либо голосовое сообщение"
        )
        bot.register_next_step_handler(message, user_skills_async)
    else:
        skills[message.chat.id] = message.text
        await end(message.chat.id)
        return

user_skills_async = async_handler(user_skills)

def create_restart_menu():
    """Создает меню с кнопками рестарта и главного меню.

    Returns:
        types.ReplyKeyboardMarkup: Объект клавиатуры с кнопками
    """
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    buttons = [
        types.KeyboardButton("🔄 Рестарт"),
        types.KeyboardButton("🏠 Главное меню")
    ]
    markup.add(*buttons)
    return markup

def add_restart_menu(message):
    """Добавляет кнопки рестарта и главного меню.
    
    Args:
        message (types.Message): Объект сообщения от пользователя
    """
    bot.send_message(
        message.chat.id, 
        "Выберите действие:", 
        reply_markup=create_restart_menu()
    )

async def end(user_id):
    """Завершает процесс создания резюме и отправляет результат пользователю."""
    try:
        bot.send_message(user_id, 'Создаю резюме...')
        _proj = ''
        for i in range(len(projects[user_id])):
            _comp = await compile(projects[user_id][i], user_id)
            _res_proj = await resume_proj(_comp, user_id)
            _proj += "\n\n" + _res_proj

        _proj = _proj.split("\n\n")

        resume_file = create_resume(
            name[user_id],
            'phone',
            'age',
            'email',
            'education',
            _proj,
            skills[user_id],
            achievements[user_id],
            'additional_info'
        )
        bot.send_document(
            chat_id=user_id,
            document=resume_file,
            visible_file_name=f"{name[user_id]}_Резюме.docx"
        )
        
        # Отправляем сообщение с кнопками рестарта и главного меню
        markup = create_restart_menu()
        bot.send_message(
            user_id,
            "Резюме готово! Выберите действие:",
            reply_markup=markup
        )
        
    except Exception as e:
        log.error(f"Error in end function: {str(e)}")
        bot.send_message(user_id, "Произошла ошибка при создании резюме. Пожалуйста, попробуйте позже.")
        # В случае ошибки также добавляем кнопки рестарта и главного меню
        markup = create_restart_menu()
        bot.send_message(
            user_id,
            "Выберите действие:",
            reply_markup=markup
        )


async def compile(answers, chat_id):
    """Компилирует ответы о проектах в структурированный формат.

    Args:
        answers (str): Текст с ответами о проектах
        chat_id (int): ID чата пользователя

    Returns:
        str: Структурированное описание проекта
    """
    prompt_compile[chat_id] = (
        "Ты - опытный составитель резюме. Я - кандидат на должность в компанию. "
        "Вот, что я сказал в беседе с тобой о своих проектах:\n"
        f'"{answers}"\n'
        "Используя мои ответы, выдели каким проектом я занимался и опиши его. "
        "Основывайся только на том, что я сказал. Не придумывай никакую новую "
        "информацию.\n"
        "Формат вывода: три bullet-point'а, разделенных символом переноса строки"
    )
    result = await send_prompt_to_gpt(prompt_compile[chat_id])
    return result


async def ask_follow_up(question_type, dialogue, context, chat_id):
    """Генерирует дополнительный вопрос на основе предыдущего диалога.

    Args:
        question_type (str): Тип вопроса
        dialogue (str): Текст диалога
        context (str): Контекст вопроса
        chat_id (int): ID чата пользователя

    Returns:
        str: Сгенерированный вопрос
    """
    prompt_ask[chat_id] = (
        "Ты - опытный собеседующий в компанию. Я - кандидат на должность в "
        "компанию. Между нами состоялся следующий диалог:\n\n"
        f"{dialogue}\n\n{question_type}\n{context}"
    )
    result = await send_prompt_to_gpt(prompt_ask[chat_id])
    return result


async def completeness(question_type, dialogue, chat_id):
    """Оценивает полноту ответа на вопрос.

    Args:
        question_type (str): Тип вопроса
        dialogue (str): Текст диалога
        chat_id (int): ID чата пользователя

    Returns:
        str: Оценка полноты ответа
    """
    prompt_compl[chat_id] = (
        "Ты - опытный составитель резюме. Я - кандидат на должность в компанию. "
        "В процессе составления резюме между нами состоял следующий диалог:\n\n"
        f"{dialogue}\n\n{question_type}"
    )
    result = await send_prompt_to_gpt(prompt_compl[chat_id])
    return result


async def resume_proj(text, chat_id):
    """Форматирует описание проекта для резюме.

    Args:
        text (str): Текст с описанием проекта
        chat_id (int): ID чата пользователя

    Returns:
        str: Отформатированное описание проекта
    """
    prompt_resume_proj[chat_id] = (
        "Ты - опытный составитель резюме с опытом работы более 10 лет. "
        "Представь, что тебе нужно написать свое резюме, а именно ту часть, "
        "где ты рассказываешь о своих проектах. Вот твои проекты:\n"
        f"{text}\n"
        "Формат вывода: напиши от своего лица часть твоего резюме, описывающая "
        "твои проекты. Будь краток и используй формальный стиль написания."
    )
    result = await send_prompt_to_gpt(prompt_resume_proj[chat_id])
    return result


def create_resume(name, phone, age, email, education, work_experience, skills,
                  achievements, additional_info):
    """Создает документ резюме в формате DOCX.

    Args:
        name (str): ФИО кандидата
        phone (str): Номер телефона
        age (str): Возраст
        email (str): Email
        education (str): Образование
        work_experience (list): Список опыта работы
        skills (str): Навыки
        achievements (str): Достижения
        additional_info (str): Дополнительная информация

    Returns:
        BytesIO: Объект с документом резюме
    """
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    title = doc.add_paragraph(name)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_style = title.style
    title_font = title_style.font
    title_font.bold = True
    title_font.size = Pt(14)

    contact_info = doc.add_paragraph(
        f"Тел.: {phone}  {age} лет  Mail: {email}"
    )
    contact_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("ОБРАЗОВАНИЕ", style='Heading 1')
    doc.add_paragraph(education, style='List Bullet')

    doc.add_paragraph("ОПЫТ РАБОТЫ", style='Heading 1')
    for exp in work_experience:
        doc.add_paragraph(exp, style='List Bullet')

    doc.add_paragraph("НАВЫКИ", style='Heading 1')
    doc.add_paragraph(skills, style='List Bullet')

    doc.add_paragraph("ДОСТИЖЕНИЯ", style='Heading 1')
    doc.add_paragraph(achievements, style='List Bullet')

    doc.add_paragraph("ДОПОЛНИТЕЛЬНАЯ ИНФОРМАЦИЯ", style='Heading 1')
    doc.add_paragraph(additional_info, style='List Bullet')

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


# Функции для AI интервьюера
def ai_interviewer_start(start_message):
    """Начинает процесс AI-интервью.

    Args:
        start_message (types.Message): Объект сообщения от пользователя
    """
    if hasattr(start_message, 'from_user'):
        pass
    else:
        return
    user_id = start_message.from_user.id
    log.info(f"Starting AI interviewer bot for user {user_id}")

    # Сохраняем текущий режим
    current_mode[user_id] = "interviewer"

    answers[user_id] = ''
    questions[user_id] = ''
    resume[user_id] = ''
    vacancy[user_id] = ''
    flag[user_id] = -1
    clicked_flag[user_id] = 0
    current_question_index[user_id] = 0

    # Добавляем только кнопку главного меню при старте
    markup = create_main_menu_button()
    bot.send_message(
        user_id,
        "Привет! Я бот от компании <a href='https://youroffer.ru/'>YourOffer</a>, мы помогаем найти работу "
        "мечты. Давай проведем с тобой пробное собеседование, чтобы лучше подготовить тебя к реальному интервью "
        "и добавить уверенности в себе!\n\n"
        "Отправь, пожалуйста, свое резюме в виде .pdf или .doc документа или в виде текстового сообщения",
        reply_markup=markup,
        parse_mode='HTML'
    )

    bot.register_next_step_handler(start_message, ask_resume, user_id)


def ask_resume(message, user_id):
    """Запрашивает резюме у пользователя для AI-интервью.

    Args:
        message (types.Message): Объект сообщения от пользователя
        user_id (int): ID пользователя
    """
    if message.text == "🏠 Главное меню":
        return_to_main_menu(message)
        return

    if (message.content_type != 'text' and
            message.content_type != 'document'):
        bot.send_message(message.from_user.id,
                         "Пожалуйста, отправь резюме в виде текстового сообщения или документа.")
        bot.register_next_step_handler(message, ask_resume, user_id)
        return

    if message.content_type == 'text':
        resume[user_id] = message.text
    elif message.content_type == 'document' and (
            message.document.file_name.lower().endswith('.pdf') or
            message.document.file_name.lower().endswith('.doc')
    ):
        try:
            file_info = bot.get_file(message.document.file_id)
            downloaded_file = bot.download_file(file_info.file_path)

            local_file_path = os.path.join("Documents", file_info.file_path.split('/')[-1])
            os.makedirs(os.path.dirname(local_file_path), exist_ok=True)

            with open(local_file_path, 'wb') as new_file:
                new_file.write(downloaded_file)

            resume[user_id] = process_pdf(local_file_path)
            os.remove(local_file_path)
        except Exception as e:
            bot.send_message(message.from_user.id,
                             "Произошла ошибка :( Пожалуйста, вернитесь в главное меню")
            return_to_main_menu(message)
            return

    bot.send_message(user_id, "Отправь, пожалуйста, описание вакансии в текстовом формате.")
    bot.register_next_step_handler(message, ask_vacancy, user_id)


def ask_vacancy(message, user_id=None):
    if message.text == "🏠 Главное меню":
        return_to_main_menu(message)
        return

    if message.content_type != 'text':
        bot.send_message(message.chat.id, "Пожалуйста, отправь описание вакансии в виде текстового сообщения.")
        bot.register_next_step_handler(message, ask_vacancy)
        return

    vacancy[message.chat.id] = message.text
    bot.send_message(message.chat.id,
                     "Спасибо! Теперь я подготовлю для тебя вопросы на основе твоего резюме и описания вакансии.")

    # Генерация вопросов на основе резюме и вакансии
    generate_questions(message.chat.id)

    bot.send_message(message.chat.id, "Вот мои вопросы:\n\n" + questions[message.chat.id])
    bot.send_message(message.chat.id, "Пожалуйста, ответь на первый вопрос.")
    bot.register_next_step_handler(message, process_answer, message.chat.id)


def process_answer(message, user_id):
    """Обрабатывает ответ пользователя на вопрос собеседования.

    Args:
        message (types.Message): Объект сообщения от пользователя
        user_id (int): ID пользователя
    """
    log.info(f"Получен ответ от пользователя {user_id} на вопрос {current_question_index.get(user_id, 0) + 1}")
    log.info(f"Тип сообщения: {message.content_type}")
    log.info(f"Текст ответа: {message.text if message.content_type == 'text' else 'Голосовое сообщение'}")

    if message.text == "🏠 Главное меню":
        log.info(f"Пользователь {user_id} вернулся в главное меню")
        return_to_main_menu(message)
        return

    if message.content_type not in ('text', 'voice'):
        log.warning(f"Пользователь {user_id} отправил неподдерживаемый тип сообщения: {message.content_type}")
        bot.send_message(user_id, "Пожалуйста, отправь ответ в виде текстового сообщения или голосового сообщения.")
        bot.register_next_step_handler(message, process_answer, user_id)
        return

    if message.content_type == 'text':
        answer = message.text
    else:
        answer = "Голосовое сообщение получено"

    # Инициализируем словарь, если его нет
    if user_id not in answers:
        answers[user_id] = ''
        log.info(f"Инициализация словаря ответов для пользователя {user_id}")
    if user_id not in current_question_index:
        current_question_index[user_id] = 0
        log.info(f"Инициализация счетчика вопросов для пользователя {user_id}")

    # Добавляем ответ в список ответов
    answers[user_id] += f"\nВопрос {current_question_index[user_id] + 1}: {answer}"
    log.info(f"Добавлен ответ на вопрос {current_question_index[user_id] + 1} для пользователя {user_id}")

    # Увеличиваем счетчик вопросов
    current_question_index[user_id] += 1
    log.info(f"Текущий индекс вопроса для пользователя {user_id}: {current_question_index[user_id]}")

    # Если это не последний вопрос
    if current_question_index[user_id] < 3:
        log.info(f"Запрашиваем следующий вопрос ({current_question_index[user_id] + 1}) у пользователя {user_id}")
        bot.send_message(user_id, f"Спасибо! Теперь ответь на вопрос {current_question_index[user_id] + 1}.")
        bot.register_next_step_handler(message, process_answer, user_id)
    else:
        log.info(f"Все вопросы пройдены для пользователя {user_id}. Начинаем анализ.")
        # Если это последний вопрос, анализируем все ответы
        analyze_interview(user_id)
        # Сбрасываем счетчик вопросов для следующего использования
        current_question_index[user_id] = 0
        log.info(f"Сброшен счетчик вопросов для пользователя {user_id}")


def analyze_interview(user_id):
    """Анализирует ответы пользователя и формирует рекомендации.

    Args:
        user_id (int): ID пользователя
    """
    prompt = f"""Проанализируй следующие ответы кандидата на вопросы собеседования и составь рекомендации по улучшению.

    Вопросы и ответы:
    {answers[user_id]}

    Составь краткий анализ и рекомендации по улучшению ответов.
    """

    analysis = send_prompt_to_gpt_sync(prompt)

    # Отправляем анализ и рекомендации
    bot.send_message(user_id, "Спасибо за участие в собеседовании! Вот мой анализ и рекомендации:\n\n" + analysis)
    
    # Добавляем кнопки рестарта и главного меню
    markup = create_restart_menu()
    bot.send_message(
        user_id,
        "Собеседование завершено! Выберите действие:",
        reply_markup=markup
    )


# Функции для parser_bot
def parser_start(message):
    """
    Начинает процесс поиска вакансий.

    Args:
        message (types.Message): Объект сообщения от пользователя
    """
    log.info(f"Starting parser bot for user {message.from_user.id}")

    # Сохраняем текущий режим
    current_mode[message.from_user.id] = "parser"

    # Добавляем только кнопку главного меню при старте
    markup = create_main_menu_button()
    bot.send_message(
        message.chat.id,
        "Привет! Я помогу тебе найти подходящие вакансии.\n"
        "Введите ключевые слова для поиска (например: 'python developer' или 'data scientist'):",
        reply_markup=markup
    )
    bot.register_next_step_handler(message, process_search_query)


def process_search_query(message):
    """
    Обрабатывает поисковый запрос и ищет вакансии.

    Args:
        message (types.Message): Объект сообщения от пользователя
    """
    if message.text == "🏠 Главное меню":
        return_to_main_menu(message)
        return

    log.info(f"User {message.from_user.id} searching for: {message.text}")

    try:
        # Формируем URL для поиска вакансий
        search_query = message.text.replace(' ', '+')
        url = f"https://api.hh.ru/vacancies?text={search_query}&per_page=5"

        # Добавляем заголовки для имитации браузера
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }

        # Выполняем запрос
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        data = response.json()

        if 'items' in data and data['items']:
            bot.send_message(message.chat.id, f"Найдено {len(data['items'])} вакансий. Показываю первые 5:")

            for vacancy in data['items']:
                # Форматируем зарплату
                salary = vacancy.get('salary', {})
                salary_text = ""
                if salary:
                    if salary.get('from') and salary.get('to'):
                        salary_text = f"от {salary['from']} до {salary['to']} {salary.get('currency', '')}"
                    elif salary.get('from'):
                        salary_text = f"от {salary['from']} {salary.get('currency', '')}"
                    elif salary.get('to'):
                        salary_text = f"до {salary['to']} {salary.get('currency', '')}"

                # Форматируем описание
                description = vacancy.get('snippet', {}).get('requirement', '')
                if description:
                    description = description[:200] + "..." if len(description) > 200 else description

                # Формируем сообщение
                vacancy_text = (
                    f"🔹 {vacancy.get('name', 'Название не указано')}\n"
                    f"💰 {salary_text if salary_text else 'Зарплата не указана'}\n"
                    f"🏢 {vacancy.get('employer', {}).get('name', 'Компания не указана')}\n"
                    f"📍 {vacancy.get('area', {}).get('name', 'Город не указан')}\n"
                    f"💼 {vacancy.get('schedule', {}).get('name', 'Формат работы не указан')}\n\n"
                    f"📝 {description}\n\n"
                    f"🔗 https://hh.ru/vacancy/{vacancy.get('id')}"
                )

                bot.send_message(message.chat.id, vacancy_text)
        else:
            bot.send_message(message.chat.id, "К сожалению, по вашему запросу ничего не найдено.")

    except requests.exceptions.RequestException as e:
        log.error(f"Error in vacancy search: {e}")
        bot.send_message(message.chat.id, "Произошла ошибка при поиске вакансий. Пожалуйста, попробуйте позже.")
    except Exception as e:
        log.error(f"Unexpected error in vacancy search: {e}")
        bot.send_message(message.chat.id, "Произошла непредвиденная ошибка. Пожалуйста, попробуйте позже.")

    # После завершения поиска добавляем кнопки рестарта и главного меню
    markup = create_restart_menu()
    bot.send_message(
        message.chat.id,
        "Поиск завершен! Выберите действие:",
        reply_markup=markup
    )

def restart_cover_letter(message):
    """Перезапускает режим создания сопроводительного письма."""
    try:
        log.info(f"Пользователь {message.from_user.id} начал перезапуск режима сопроводительного письма")
        
        # Очищаем все данные пользователя
        log.info(f"Очистка данных пользователя {message.from_user.id} в режиме сопроводительного письма")
        if message.chat.id in resume:
            del resume[message.chat.id]
            log.debug(f"Удалены данные резюме для пользователя {message.from_user.id}")
        if message.chat.id in vacancy:
            del vacancy[message.chat.id]
            log.debug(f"Удалены данные вакансии для пользователя {message.from_user.id}")
        if message.chat.id in questions:
            del questions[message.chat.id]
            log.debug(f"Удалены вопросы для пользователя {message.from_user.id}")
        if message.chat.id in answers:
            del answers[message.chat.id]
            log.debug(f"Удалены ответы для пользователя {message.from_user.id}")
        if message.chat.id in current_question_index:
            del current_question_index[message.chat.id]
            log.debug(f"Удален индекс вопросов для пользователя {message.from_user.id}")
        
        # Перезапускаем режим с сохранением текущего режима
        current_mode[message.from_user.id] = "cover_letter"
        log.info(f"Установлен режим 'cover_letter' для пользователя {message.from_user.id}")
        
        # Инициализируем необходимые переменные
        answers[message.chat.id] = ''
        current_question_index[message.chat.id] = 0
        log.debug(f"Инициализированы новые переменные для пользователя {message.from_user.id}")
        
        # Отправляем приветственное сообщение с кнопкой главного меню
        markup = create_main_menu_button()
        log.info(f"Отправка приветственного сообщения пользователю {message.from_user.id}")
        bot.send_message(
            message.chat.id,
            "Давайте начнем заново!\n\n"
            "Отправь, пожалуйста, свое резюме в текстовом формате, или в форматах pdf или doc",
            reply_markup=markup
        )
        bot.register_next_step_handler(message, ask_resume_async)
        log.info(f"Установлен обработчик следующего шага для пользователя {message.from_user.id}")
        
    except Exception as e:
        log.error(f"Ошибка при перезапуске режима сопроводительного письма для пользователя {message.from_user.id}: {str(e)}")
        # В случае ошибки возвращаем в главное меню
        return_to_main_menu(message)

def restart_resume_bot(message):
    """Перезапускает режим создания резюме."""
    try:
        log.info(f"Пользователь {message.from_user.id} начал перезапуск режима создания резюме")
        
        # Очищаем все данные пользователя
        log.info(f"Очистка данных пользователя {message.from_user.id} в режиме создания резюме")
        if message.chat.id in name:
            del name[message.chat.id]
            log.debug(f"Удалено имя пользователя {message.from_user.id}")
        if message.chat.id in summary:
            del summary[message.chat.id]
            log.debug(f"Удалено описание пользователя {message.from_user.id}")
        if message.chat.id in projects:
            del projects[message.chat.id]
            log.debug(f"Удалены проекты пользователя {message.from_user.id}")
        if message.chat.id in skills:
            del skills[message.chat.id]
            log.debug(f"Удалены навыки пользователя {message.from_user.id}")
        if message.chat.id in achievements:
            del achievements[message.chat.id]
            log.debug(f"Удалены достижения пользователя {message.from_user.id}")
        if message.chat.id in answers_X:
            del answers_X[message.chat.id]
            log.debug(f"Удалены ответы X пользователя {message.from_user.id}")
        if message.chat.id in answers_Y:
            del answers_Y[message.chat.id]
            log.debug(f"Удалены ответы Y пользователя {message.from_user.id}")
        if message.chat.id in answers_Z:
            del answers_Z[message.chat.id]
            log.debug(f"Удалены ответы Z пользователя {message.from_user.id}")
        if message.chat.id in dialogue:
            del dialogue[message.chat.id]
            log.debug(f"Удален диалог пользователя {message.from_user.id}")
        if message.chat.id in context:
            del context[message.chat.id]
            log.debug(f"Удален контекст пользователя {message.from_user.id}")
        if message.chat.id in question_counter:
            del question_counter[message.chat.id]
            log.debug(f"Удален счетчик вопросов пользователя {message.from_user.id}")
        if message.chat.id in follow_up:
            del follow_up[message.chat.id]
            log.debug(f"Удалены follow-up вопросы пользователя {message.from_user.id}")
        if message.chat.id in prompt_compile:
            del prompt_compile[message.chat.id]
            log.debug(f"Удалены скомпилированные промпты пользователя {message.from_user.id}")
        if message.chat.id in prompt_ask:
            del prompt_ask[message.chat.id]
            log.debug(f"Удалены промпты вопросов пользователя {message.from_user.id}")
        if message.chat.id in prompt_compl:
            del prompt_compl[message.chat.id]
            log.debug(f"Удалены промпты полноты пользователя {message.from_user.id}")
        if message.chat.id in prompt_resume_proj:
            del prompt_resume_proj[message.chat.id]
            log.debug(f"Удалены промпты проектов пользователя {message.from_user.id}")
        if message.chat.id in previous_message_id:
            del previous_message_id[message.chat.id]
            log.debug(f"Удален ID предыдущего сообщения пользователя {message.from_user.id}")
        
        # Перезапускаем режим с сохранением текущего режима
        current_mode[message.from_user.id] = "resume"
        log.info(f"Установлен режим 'resume' для пользователя {message.from_user.id}")
        
        # Инициализируем необходимые переменные
        answers_X[message.chat.id] = ''
        answers_Y[message.chat.id] = ''
        answers_Z[message.chat.id] = ''
        dialogue[message.chat.id] = ''
        projects[message.chat.id] = []
        context[message.chat.id] = []
        question_counter[message.chat.id] = 1
        log.debug(f"Инициализированы новые переменные для пользователя {message.from_user.id}")
        
        # Отправляем приветственное сообщение с кнопкой главного меню
        markup = create_main_menu_button()
        log.info(f"Отправка приветственного сообщения пользователю {message.from_user.id}")
        bot.send_message(
            message.chat.id,
            "Давайте начнем заново!\n\n"
            "Напиши, пожалуйста, свое ФИО.",
            reply_markup=markup
        )
        bot.register_next_step_handler(message, user_name)
        log.info(f"Установлен обработчик следующего шага для пользователя {message.from_user.id}")
        
    except Exception as e:
        log.error(f"Ошибка при перезапуске режима создания резюме для пользователя {message.from_user.id}: {str(e)}")
        # В случае ошибки возвращаем в главное меню
        return_to_main_menu(message)

def restart_ai_interviewer(message):
    """Перезапускает режим AI-интервьюера."""
    try:
        log.info(f"Пользователь {message.from_user.id} начал перезапуск режима AI-интервьюера")
        
        # Очищаем все данные пользователя
        log.info(f"Очистка данных пользователя {message.from_user.id} в режиме AI-интервьюера")
        if message.chat.id in answers:
            del answers[message.chat.id]
            log.debug(f"Удалены ответы пользователя {message.from_user.id}")
        if message.chat.id in questions:
            del questions[message.chat.id]
            log.debug(f"Удалены вопросы пользователя {message.from_user.id}")
        if message.chat.id in resume:
            del resume[message.chat.id]
            log.debug(f"Удалено резюме пользователя {message.from_user.id}")
        if message.chat.id in vacancy:
            del vacancy[message.chat.id]
            log.debug(f"Удалена вакансия пользователя {message.from_user.id}")
        if message.chat.id in flag:
            del flag[message.chat.id]
            log.debug(f"Удален флаг пользователя {message.from_user.id}")
        if message.chat.id in clicked_flag:
            del clicked_flag[message.chat.id]
            log.debug(f"Удален clicked_flag пользователя {message.from_user.id}")
        if message.chat.id in current_question_index:
            del current_question_index[message.chat.id]
            log.debug(f"Удален индекс вопросов пользователя {message.from_user.id}")
        
        # Перезапускаем режим с сохранением текущего режима
        current_mode[message.from_user.id] = "interviewer"
        log.info(f"Установлен режим 'interviewer' для пользователя {message.from_user.id}")
        
        # Инициализируем необходимые переменные
        answers[message.chat.id] = ''
        questions[message.chat.id] = ''
        current_question_index[message.chat.id] = 0
        flag[message.chat.id] = -1
        clicked_flag[message.chat.id] = 0
        log.debug(f"Инициализированы новые переменные для пользователя {message.from_user.id}")
        
        # Отправляем приветственное сообщение с кнопкой главного меню
        markup = create_main_menu_button()
        log.info(f"Отправка приветственного сообщения пользователю {message.from_user.id}")
        bot.send_message(
            message.chat.id,
            "Давайте начнем собеседование заново!\n\n"
            "Отправь, пожалуйста, свое резюме в виде .pdf или .doc документа или в виде текстового сообщения",
            reply_markup=markup,
            parse_mode='HTML'
        )
        bot.register_next_step_handler(message, ask_resume, message.chat.id)
        log.info(f"Установлен обработчик следующего шага для пользователя {message.from_user.id}")
        
    except Exception as e:
        log.error(f"Ошибка при перезапуске режима AI-интервьюера для пользователя {message.from_user.id}: {str(e)}")
        # В случае ошибки возвращаем в главное меню
        return_to_main_menu(message)

def restart_parser(message):
    """Перезапускает режим поиска вакансий."""
    try:
        log.info(f"Пользователь {message.from_user.id} начал перезапуск режима парсера вакансий")
        
        # Перезапускаем режим с сохранением текущего режима
        current_mode[message.from_user.id] = "parser"
        log.info(f"Установлен режим 'parser' для пользователя {message.from_user.id}")
        
        # Отправляем приветственное сообщение с кнопкой главного меню
        markup = create_main_menu_button()
        log.info(f"Отправка приветственного сообщения пользователю {message.from_user.id}")
        bot.send_message(
            message.chat.id,
            "Давайте начнем поиск заново!\n\n"
            "Введите ключевые слова для поиска (например: 'python developer' или 'data scientist'):",
            reply_markup=markup
        )
        bot.register_next_step_handler(message, process_search_query)
        log.info(f"Установлен обработчик следующего шага для пользователя {message.from_user.id}")
        
    except Exception as e:
        log.error(f"Ошибка при перезапуске режима парсера вакансий для пользователя {message.from_user.id}: {str(e)}")
        # В случае ошибки возвращаем в главное меню
        return_to_main_menu(message)

def generate_questions(user_id):
    """Генерирует вопросы для собеседования на основе резюме и описания вакансии.

    Args:
        user_id (int): ID пользователя
    """
    log.info(f"Генерация вопросов для пользователя {user_id}")
    log.info(f"Резюме пользователя: {resume[user_id][:100]}...")  # Логируем первые 100 символов резюме
    log.info(f"Описание вакансии: {vacancy[user_id][:100]}...")  # Логируем первые 100 символов вакансии

    prompt = f"""На основе следующего резюме и описания вакансии составь 3 четких и конкретных вопроса для собеседования.
    Вопросы должны быть пронумерованы от 1 до 3.
    Каждый вопрос должен быть на новой строке.
    Вопросы должны быть направлены на оценку соответствия кандидата требованиям вакансии.

    Резюме:
    {resume[user_id]}

    Вакансия:
    {vacancy[user_id]}

    Формат вывода:
    1. Первый вопрос
    2. Второй вопрос
    3. Третий вопрос
    """

    questions[user_id] = send_prompt_to_gpt_sync(prompt)
    log.info(f"Сгенерированные вопросы для пользователя {user_id}: {questions[user_id]}")

def send_prompt_to_gpt_sync(prompt):
    """Синхронная версия функции отправки запроса к GPT API.

    Args:
        prompt (str): Текст запроса к GPT

    Returns:
        str: Ответ от GPT или сообщение об ошибке
    """
    endpoint = 'https://api.openai.com/v1/chat/completions'

    headers = {
        'Content-Type': 'application/json',
        'Authorization': f'Bearer {api_key}'
    }

    data = {
        'model': 'gpt-4o-mini',
        'messages': [{'role': 'user', 'content': prompt}],
        'max_tokens': 3000,
        'top_p': 1.0,
        'temperature': 0.6
    }

    response = requests.post(endpoint, headers=headers, json=data)

    if response.status_code == 200:
        response_data = json.loads(response.text)
        return response_data['choices'][0]['message']['content']
    else:
        print("Error:", response.text)
        return "Извините, произошла ошибка при генерации вопросов."